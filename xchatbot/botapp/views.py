from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from .models import Bot
import os
import uuid
import time
import traceback
import logging
from groq import Groq

# Setup logger
logger = logging.getLogger(__name__)

GROQ_API_KEY = "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"

try:
    client = Groq(api_key=GROQ_API_KEY)
except Exception as e:
    logger.exception("Failed to initialize Groq client")
    raise

@login_required
def character_prompt(request):
    try:
        if request.method == "POST":
            name = request.POST.get("name", "").strip()
            system_prompt = request.POST.get("system_prompt", "").strip()
            
            if not name or not system_prompt:
                return JsonResponse({"error": "Name and system prompt required"}, status=400)
            
            bot = Bot.objects.create(
                owner=request.user,
                name=name,
                system_prompt=system_prompt
            )
            return redirect('public_chat_ui', public_id=bot.public_id)
        
        return render(request, 'mainapp/character_prompt.html')
    except Exception as e:
        logger.exception("Error in character_prompt view")
        return JsonResponse({"error": str(e)}, status=500)

def public_chat_ui(request, public_id):
    try:
        bot = get_object_or_404(Bot, public_id=public_id)
        session_key = f'chat_history_{public_id}'
        return render(request, 'mainapp/chat_ui.html', {
            'bot': bot,
            'chat': request.session.get(session_key, [])
        })
    except Exception as e:
        logger.exception("Error in public_chat_ui view")
        return JsonResponse({"error": str(e)}, status=500)
    
from .models import BotUsages

def send_message_public(request, public_id):
    try:
        bot = get_object_or_404(Bot, public_id=public_id)
        user_message = request.POST.get("message", "").strip()
        
        if not user_message:
            return JsonResponse({"error": "Message is required"}, status=400)
        
        session_key = f'chat_history_{public_id}'
        history = request.session.get(session_key, [])[-5:]
        
        # Enhanced system prompt with knowledge base integration
        system_content = f"""{bot.system_prompt}

You have access to the following knowledge base. Use this information to answer user questions naturally and intelligently, just like a human expert would.

KNOWLEDGE BASE:
{bot.knowledge_data}

Instructions:
- You are an intelligent AI assistant with access to comprehensive information above
- Answer questions naturally using your understanding of the knowledge base
- Think about what the user is really asking and provide helpful, relevant responses
- Be conversational, analytical, and genuinely helpful
- When listing information, use line breaks between items for better readability

NEVER use ASCII emoticons like (；´Д｀) (´；） (≖ω≖) ಠ_ಠ or any text-based faces
Only use real Unicode emojis if needed: 😊 😂 ❤️ 👍"""

        messages = [{"role": "system", "content": system_content}]
        
        for line in history:
            messages.append({"role": "user", "content": line['user']})
            messages.append({"role": "assistant", "content": line['ai']})
        
        messages.append({"role": "user", "content": user_message})
        
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=messages,
            max_tokens=400,
            temperature=0.6,
            top_p=0.9,
            stop=["Human:", "User:", "\n\nUser:", "\n\nHuman:"],
            frequency_penalty=0.0,
            presence_penalty=0.0
        )
        
        ai_reply = response.choices[0].message.content.strip()

        words_count = len(user_message.split())
        messages_count = 1

        usage, created = BotUsages.objects.get_or_create(bot=bot)
        usage.bot_messages = str(int(usage.bot_messages or 0) + messages_count)
        usage.bot_words = str(int(usage.bot_words or 0) + words_count)
        usage.save()
        
        # General formatting for ANY type of agent - no hardcoded values
        import re
        
        # Remove markdown formatting
        ai_reply = ai_reply.replace('**', '')
        ai_reply = ai_reply.replace('* ', '')
        ai_reply = ai_reply.replace('• ', '')
        
        # General pattern: Add line break before any word followed by colon and price/number
        ai_reply = re.sub(r'([A-Za-z0-9\s]+):\s*(₹|Rs|\$|[0-9])', r'\n\n\1: \2', ai_reply)
        
        # Add line break before common section headers (any word ending with "ing" or containing "Plan" etc)
        ai_reply = re.sub(r'([A-Za-z]+(?:ing|Plans?|Services?|Options?)[:\s]*)', r'\n\n\1', ai_reply)
        
        # Clean up multiple line breaks
        ai_reply = re.sub(r'\n{3,}', '\n\n', ai_reply)
        ai_reply = ai_reply.strip()
        
        # Remove ASCII emoticons and kaomoji
        ai_reply = re.sub(r'\([；´Д｀\)]+\)', '', ai_reply)
        ai_reply = re.sub(r'\([^)]*[；´Д｀ಠ≖ω][^)]*\)', '', ai_reply)
        ai_reply = re.sub(r'[（）\(\)][^）\)]*[；´Д｀ಠ≖ω][^）\)]*[（）\(\)]', '', ai_reply)
        ai_reply = ai_reply.strip()
        
        if not ai_reply:
            raise ValueError("Empty response from model")
        
        history.append({"user": user_message, "ai": ai_reply})
        request.session[session_key] = history
        
        return JsonResponse({"reply": ai_reply})
        
    except Exception as e:
        logger.exception("Error in send_message_public view")
        return JsonResponse({"error": str(e)}, status=500)

# Createing of bot creation ----------------------------------------------------------------------------------------
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import XBot

# @login_required
# def create_bot(request):
#     if request.method == 'POST':
#         bot_name = request.POST.get('bot_name')
#         description = request.POST.get('description', '')
#         knowledge_base = request.POST.get('knowledge_base', '')
#         website_url = request.POST.get('website_url', '')
#         pdf_file = request.FILES.get('pdf_file')
#         word_file = request.FILES.get('word_file')
        
#         if not bot_name:
#             messages.error(request, 'Bot name is required.')
#             return render(request, 'xbot/create_bot.html')
        
#         bot = XBot.objects.create(
#             owner=request.user,
#             name=bot_name,
#             description=description,
#             knowledge=knowledge_base,
#             pdf_files=pdf_file,
#             word_files=word_file,
#             website_links=website_url
#         )
        
#         messages.success(request, f'Bot "{bot_name}" created successfully!')
#         return redirect('bot_detail', public_id=bot.public_id)
    
#     return render(request, 'workflow/create_bot.html')
from django.views.decorators.csrf import csrf_exempt
@csrf_exempt
@login_required
def create_bot(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        description = request.POST.get('description')
        text = request.POST.get('text')

        # Save bot to database
        bot = Bot.objects.create(
            owner=request.user,
            name=title,
            system_prompt=description,
            knowledge_data=text
        )

        return JsonResponse({'success': True, 'bot_id': bot.id})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import os
import PyPDF2
from docx import Document
import os
from PyPDF2 import PdfReader
import docx
@csrf_exempt
def upload_file(request):
    if request.method == 'POST':
        bot_id = request.POST.get('bot_id')
        uploaded_file = request.FILES.get('file')

        if not bot_id or not uploaded_file:
            return JsonResponse({'success': False, 'error': 'Missing bot_id or file'})

        try:
            bot = Bot.objects.get(id=int(bot_id))
        except Bot.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Bot not found'})

        # Save uploaded file
        bot.data_files.save(uploaded_file.name, uploaded_file)

        # Read text based on file type
        ext = os.path.splitext(uploaded_file.name)[1].lower()
        text = ""

        try:
            if ext == ".pdf":
                reader = PdfReader(bot.data_files.path)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
            elif ext in [".doc", ".docx"]:
                doc = docx.Document(bot.data_files.path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == ".txt":
                with open(bot.data_files.path, 'r', encoding="utf-8") as f:
                    text = f.read()
            else:
                return JsonResponse({'success': False, 'error': 'Unsupported file type'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

        # Append extracted text to knowledge
        bot.knowledge_data = (bot.knowledge_data or "") + "\n" + text
        bot.save()

        return JsonResponse({'success': True})
    else:
        return JsonResponse({'success': False, 'error': 'Invalid request'})

from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from django.utils.html import strip_tags
import requests
from bs4 import BeautifulSoup
import json
from .models import Bot

@csrf_exempt
@require_POST
def train_from_url(request):
    try:
        data = json.loads(request.body)
        url = data.get("url")
        bot_id = data.get("bot_id")

        if not url or not bot_id:
            return JsonResponse({'success': False, 'error': 'Missing URL or bot ID'})

        bot = Bot.objects.get(id=bot_id)

        # Crawl and extract text
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")
        extracted_text = strip_tags(soup.get_text())

        # Append to knowledge
        bot.knowledge_data = (bot.knowledge_data or "") + "\n" + extracted_text.strip()
        bot.save()

        return JsonResponse({'success': True})
    except Bot.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Bot not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


from django.http import JsonResponse
from django.shortcuts import render, get_object_or_404
from .models import Bot
def edit_agents(request, bot_id):
    bot = get_object_or_404(Bot, id=bot_id, owner=request.user)

    usages = BotUsages.objects.filter(bot=bot) 
    
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        # Get form data
        name = request.POST.get('name', '').strip()
        system_prompt = request.POST.get('system_prompt', '').strip()
        agent_nature = request.POST.get('agent_nature', '')
        agent_vibe = request.POST.get('agent_vibe', '')
        ai_model = request.POST.get('ai_model', '')
        server_type = request.POST.get('server_type', '')
        server_location = request.POST.get('server_location', '')
        storage_size = request.POST.get('storage_size', '')
        backup_frequency = request.POST.get('backup_frequency', '')
        response_speed = request.POST.get('response_speed', '')
        conversation_memory = request.POST.get('conversation_memory', '')
        
        # Validate required fields
        if not name or not system_prompt:
            return JsonResponse({
                'success': False, 
                'message': 'Name and description are required'
            }, status=400)
        
        # Update bot fields
        try:
            bot.name = name
            bot.system_prompt = system_prompt
            
            # Update new fields if provided
            if agent_nature:
                bot.agent_nature = agent_nature
            if agent_vibe:
                bot.agent_vibe = agent_vibe
            if ai_model:
                bot.ai_model = ai_model
            if server_type:
                bot.server_type = server_type
            if server_location:
                bot.server_location = server_location
            if storage_size:
                bot.storage_size = storage_size
            if backup_frequency:
                bot.backup_frequency = backup_frequency
            if response_speed:
                bot.response_speed = response_speed
            if conversation_memory:
                bot.conversation_memory = conversation_memory
            
            bot.save()
            
            return JsonResponse({
                'success': True, 
                'message': 'Bot updated successfully'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False, 
                'message': f'Error updating bot: {str(e)}'
            }, status=500)
    
    return render(request, 'workflow/edit_bot.html', {'bot': bot, 'usages': usages})



import os
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from .models import Bot

from PyPDF2 import PdfReader
import docx
@csrf_exempt
def edit_upload_file(request, bot_id):
    if request.method == 'POST':
        bot = get_object_or_404(Bot, id=bot_id)
        files = request.FILES.getlist('files[]')
        manual_text = request.POST.get('text', '').strip()

        combined_text = ""

        # Process files
        for file in files:
            ext = os.path.splitext(file.name)[1].lower()
            try:
                if ext == ".pdf":
                    reader = PdfReader(file)
                    text = "\n".join([page.extract_text() or "" for page in reader.pages])
                elif ext in [".doc", ".docx"]:
                    document = docx.Document(file)
                    text = "\n".join([para.text for para in document.paragraphs])
                elif ext == ".txt":
                    text = file.read().decode('utf-8', errors='ignore')
                else:
                    return JsonResponse({'success': False, 'message': f'Unsupported file type: {ext}'})
                combined_text += "\n\n" + text.strip()
            except Exception as e:
                return JsonResponse({'success': False, 'message': f'Error processing {file.name}: {str(e)}'})

        # Add manual text from textarea
        if manual_text:
            combined_text += "\n\n" + manual_text

        # Save to knowledge_data
        if combined_text:
            if bot.knowledge_data:
                bot.knowledge_data += "\n\n" + combined_text
            else:
                bot.knowledge_data = combined_text
            bot.save()
            return JsonResponse({'success': True, 'message': 'Text and files saved successfully.'})
        else:
            return JsonResponse({'success': False, 'message': 'No text or files provided.'})

    return JsonResponse({'success': False, 'message': 'Invalid request method.'})

@csrf_exempt
def fetch_website_data(request, bot_id):
    if request.method == 'POST':
        url = request.POST.get('url')

        if not url:
            return JsonResponse({'success': False, 'message': 'No URL provided.'})

        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')

            # Remove scripts and styles
            for tag in soup(['script', 'style', 'noscript']):
                tag.decompose()

            text = soup.get_text(separator='\n', strip=True)

            bot = get_object_or_404(Bot, id=bot_id)
            if bot.knowledge_data:
                bot.knowledge_data += '\n\n' + text
            else:
                bot.knowledge_data = text
            bot.save()

            return JsonResponse({'success': True, 'message': 'Website content fetched and saved successfully.'})

        except requests.RequestException as e:
            return JsonResponse({'success': False, 'message': f'Error fetching URL: {str(e)}'})

    return JsonResponse({'success': False, 'message': 'Invalid request method.'})


@csrf_exempt
def toggle_bot_status(request, bot_id):
    if request.method == 'POST':
        bot = Bot.objects.get(id=bot_id)
        bot.bot_status = not bot.bot_status
        bot.save()
        return JsonResponse({"new_status": bot.bot_status})
    

def delete_bot(request, bot_id):
    bot = get_object_or_404(Bot, id=bot_id)
    bot.delete()
    return redirect("mainapp:agents")  # use your home view name

